import docx

def extract_text_from_docx(path):
    doc = docx.Document(path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    # Replace newlines within cell to avoid breaking the row formatting
                    cleaned_cell = cell.text.strip().replace('\n', ' ')
                    # Avoid duplicate text from merged cells
                    if not row_text or row_text[-1] != cleaned_cell:
                        row_text.append(cleaned_cell)
            if row_text:
                full_text.append(" | ".join(row_text))
                
    with open("proposal_text.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(full_text))

if __name__ == "__main__":
    extract_text_from_docx(r'C:\Users\USERAS\Desktop\TGD1 papers\EP4DIS_Research_Project_Proposal_Shawon_Mehrab_March_2026_final copy.docx')
