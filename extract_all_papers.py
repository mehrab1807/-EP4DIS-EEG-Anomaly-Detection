import os
import docx
import PyPDF2
from pathlib import Path

def extract_from_docx(path, output_path):
    try:
        doc = docx.Document(path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        cleaned_cell = cell.text.strip().replace('\n', ' ')
                        if not row_text or row_text[-1] != cleaned_cell:
                            row_text.append(cleaned_cell)
                if row_text:
                    full_text.append(" | ".join(row_text))
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(full_text))
        print(f"Extracted DOCX: {path.name}")
    except Exception as e:
        print(f"Error extracting {path.name}: {e}")

def extract_from_pdf(path, output_path):
    try:
        with open(path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            full_text = []
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text = page.extract_text()
                if text:
                    full_text.append(text)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write("\n".join(full_text))
        print(f"Extracted PDF: {path.name}")
    except Exception as e:
        print(f"Error extracting {path.name}: {e}")

def main():
    source_dir = Path(r'C:\Users\USERAS\Desktop\TGD1 papers')
    dest_dir = Path(r'C:\Users\USERAS\.gemini\antigravity\scratch\eeg_anomaly_detection\extracted_papers')
    
    dest_dir.mkdir(parents=True, exist_ok=True)
    
    for file_path in source_dir.iterdir():
        if file_path.name.startswith('~'):
            continue
        
        output_name = file_path.stem + '.txt'
        output_path = dest_dir / output_name
        
        if file_path.suffix.lower() == '.docx':
            extract_from_docx(file_path, output_path)
        elif file_path.suffix.lower() == '.pdf':
            extract_from_pdf(file_path, output_path)

if __name__ == "__main__":
    main()
